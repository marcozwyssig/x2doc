import logging
import xml.etree.ElementTree as ET
from abc import ABC, abstractmethod
from typing import List, Tuple, Union, Optional

from docx import Document as DocxDocument
from docx.document import Document as DocxDocument2
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.table import _Cell, Table as DocxTable
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.shared import Cm

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(name)s: %(message)s',
    handlers=[
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Constants
TOTAL_TABLE_WIDTH_CM = 15  # Total table width in cm for Word documents


def iter_block_items(parent: Union[DocxDocument2, _Cell]):
    """
    Iterate over block-level items (tables and paragraphs) in the document.
    
    Args:
        parent: The parent document or cell from which to iterate.

    Yields:
        DocxParagraph or DocxTable objects.
    """
    logger.debug(f"Iterating block items for parent: {type(parent).__name__}")
    if isinstance(parent, DocxDocument2):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        logger.error("Invalid parent type provided to iter_block_items.")
        raise ValueError("Invalid parent type provided to iter_block_items.")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield DocxParagraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield DocxTable(child, parent)
        else:
            logger.warning(f"Unknown child type encountered: {type(child)}")


class DocumentElement(ABC):
    """
    Abstract base class for document elements.
    """

    @classmethod
    @abstractmethod
    def from_xml(cls, element: ET.Element) -> 'DocumentElement':
        """
        Create an instance from an XML element.
        """
        pass

    @classmethod
    @abstractmethod
    def from_word(cls, element: Union[DocxParagraph, DocxTable]) -> 'DocumentElement':
        """
        Create an instance from a Word document element.
        """
        pass

    @abstractmethod
    def to_word(self, docx_document: DocxDocument, level: int = 1) -> None:
        """
        Add the element to a Word document.
        """
        pass

    @abstractmethod
    def to_xml(self) -> ET.Element:
        """
        Convert the element to an XML element.
        """
        pass


class Table(DocumentElement):
    """
    Represents a table in the document.
    """

    def __init__(self, columns: List[Tuple[str, Optional[str]]], rows: List[List[str]]):
        self.columns = columns  # List of tuples (column name, width percentage)
        self.rows = rows
        logger.debug(f"Initialized Table with columns: {self.columns} and rows: {self.rows}")

    @classmethod
    def from_xml(cls, table_element: ET.Element) -> 'Table':
        """
        Create a Table instance from an XML element.
        """
        logger.info("Parsing Table from XML.")
        columns = [
            (col.text, col.attrib.get('width'))
            for col in table_element.find('columns').findall('column')
        ]
        rows = []
        rows_element = table_element.find('rows')
        if rows_element is not None:
            for row in rows_element.findall('row'):
                row_data = [cell.text or "" for cell in row.findall('cell')]
                rows.append(row_data)
        logger.debug(f"Parsed Table columns: {columns}, rows: {rows}")
        return cls(columns, rows)

    @classmethod
    def from_word(cls, table: DocxTable) -> 'Table':
        """
        Create a Table instance from a Word document table.
        """
        logger.info("Parsing Table from Word document.")
        columns = []
        rows = []
        if table.rows:
            header_cells = table.rows[0].cells
            for cell in header_cells:
                columns.append((cell.text.strip(), None))  # Width handling can be enhanced
            for row in table.rows[1:]:
                rows.append([cell.text.strip() for cell in row.cells])
        logger.debug(f"Parsed Table from Word with columns: {columns}, rows: {rows}")
        return cls(columns, rows)

    def to_word(self, docx_document: DocxDocument, level: int = 1) -> None:
        """
        Add the table to a Word document.
        """
        logger.info("Adding Table to Word document.")
        if not self.columns:
            logger.warning("No columns defined for the table.")
            return

        table = docx_document.add_table(rows=1, cols=len(self.columns))
        hdr_cells = table.rows[0].cells

        # Set column headers and widths
        for i, (col, width) in enumerate(self.columns):
            hdr_cells[i].text = col
            if width:
                try:
                    width_percentage = float(width)
                    column_width_cm = (width_percentage / 100) * TOTAL_TABLE_WIDTH_CM
                    hdr_cells[i].width = Cm(column_width_cm)
                    logger.debug(f"Set width for column '{col}' to {column_width_cm} cm.")
                except ValueError:
                    logger.warning(f"Invalid width value '{width}' for column '{col}'.")

        # Add data rows
        for row_data in self.rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                if i < len(row_cells):
                    row_cells[i].text = cell_data
        logger.debug("Table added to Word document successfully.")

    def to_xml(self) -> ET.Element:
        """
        Convert the table to an XML element.
        """
        logger.info("Converting Table to XML.")
        table_element = ET.Element('table')
        columns_element = ET.SubElement(table_element, 'columns')
        for col, width in self.columns:
            col_attrib = {'width': width} if width else {}
            col_element = ET.SubElement(columns_element, 'column', **col_attrib)
            col_element.text = col

        rows_element = ET.SubElement(table_element, 'rows')
        for row in self.rows:
            row_element = ET.SubElement(rows_element, 'row')
            for cell in row:
                cell_element = ET.SubElement(row_element, 'cell')
                cell_element.text = cell

        logger.debug("Table converted to XML successfully.")
        return table_element

    def __repr__(self) -> str:
        return f"Table(columns={self.columns}, rows={self.rows})"


class Paragraph(DocumentElement):
    """
    Represents a paragraph in the document.
    """

    def __init__(self, text: str):
        self.text = text.strip()
        logger.debug(f"Initialized Paragraph with text: {self.text}")

    @classmethod
    def from_xml(cls, paragraph_element: ET.Element) -> 'Paragraph':
        """
        Create a Paragraph instance from an XML element.
        """
        logger.info("Parsing Paragraph from XML.")
        text = paragraph_element.text or ""
        logger.debug(f"Parsed Paragraph text: {text}")
        return cls(text)

    @classmethod
    def from_word(cls, paragraph: DocxParagraph) -> 'Paragraph':
        """
        Create a Paragraph instance from a Word document paragraph.
        """
        logger.info("Parsing Paragraph from Word document.")
        text = paragraph.text
        logger.debug(f"Parsed Paragraph text: {text}")
        return cls(text)

    def to_word(self, docx_document: DocxDocument, level: int = 1) -> None:
        """
        Add the paragraph to a Word document.
        """
        logger.info("Adding Paragraph to Word document.")
        docx_document.add_paragraph(self.text)
        logger.debug("Paragraph added to Word document successfully.")

    def to_xml(self) -> ET.Element:
        """
        Convert the paragraph to an XML element.
        """
        logger.info("Converting Paragraph to XML.")
        paragraph_element = ET.Element('paragraph')
        paragraph_element.text = self.text
        logger.debug("Paragraph converted to XML successfully.")
        return paragraph_element

    def __repr__(self) -> str:
        return f"Paragraph(text={self.text})"


class Chapter(DocumentElement):
    """
    Represents a chapter in the document, which can contain subchapters, tables, and paragraphs.
    """

    def __init__(self, title: str, id_: str, elements: Optional[List[DocumentElement]] = None):
        self.title = title.strip()
        self.id_ = id_
        self.elements = elements or []
        logger.debug(f"Initialized Chapter with title: '{self.title}', id: '{self.id_}', elements count: {len(self.elements)}")

    @classmethod
    def from_xml(cls, chapter_element: ET.Element) -> 'Chapter':
        """
        Create a Chapter instance from an XML element.
        """
        logger.info("Parsing Chapter from XML.")
        title = chapter_element.attrib.get('title', 'Untitled Chapter')
        id_ = chapter_element.attrib.get('id', 'unknown-id')
        elements: List[DocumentElement] = []

        for child in chapter_element:
            if child.tag == 'chapter':
                elements.append(Chapter.from_xml(child))
            elif child.tag == 'table':
                elements.append(Table.from_xml(child))
            elif child.tag == 'paragraph':
                elements.append(Paragraph.from_xml(child))
            else:
                logger.warning(f"Unknown element '{child.tag}' encountered in Chapter.")

        logger.debug(f"Parsed Chapter '{title}' with {len(elements)} elements.")
        return cls(title, id_, elements)

    @classmethod
    def from_word(
        cls,
        paragraphs: List[Union[DocxParagraph, DocxTable]],
        current_index: int,
        heading_level: int,
        docx_document: DocxDocument
    ) -> Tuple[List['Chapter'], int]:
        """
        Recursively parse chapters from the given paragraphs starting at current_index.

        Args:
            paragraphs: List of paragraphs and tables from the Word document.
            current_index: Current index in the paragraphs list.
            heading_level: Current heading level to parse.
            docx_document: The Word document object.

        Returns:
            A tuple containing the list of parsed chapters and the updated current_index.
        """
        logger.info(f"Parsing chapters starting at index {current_index} with heading level {heading_level}.")
        chapters: List[Chapter] = []
        current_chapter: Optional[Chapter] = None

        while current_index < len(paragraphs):
            block = paragraphs[current_index]

            if isinstance(block, DocxParagraph) and block.style.name.startswith('Heading'):
                try:
                    current_heading_level = int(block.style.name.split()[-1])
                except (IndexError, ValueError):
                    logger.error(f"Invalid heading style format: '{block.style.name}'")
                    current_index += 1
                    continue

                logger.debug(f"Found heading at index {current_index} with level {current_heading_level}.")

                if current_heading_level < heading_level:
                    logger.debug("Encountered a higher-level heading. Ending current chapter parsing.")
                    break
                elif current_heading_level == heading_level:
                    if current_chapter is not None:
                        chapters.append(current_chapter)
                        logger.debug(f"Added Chapter '{current_chapter.title}' to chapters list.")
                    # Create a new chapter
                    current_chapter = Chapter(block.text, id_=f"chapter-{current_index}")
                    logger.debug(f"Started new Chapter '{current_chapter.title}'.")
                    current_index += 1
                else:  # current_heading_level > heading_level
                    # Handle subchapters recursively
                    subchapters, current_index = Chapter.from_word(
                        paragraphs, current_index, current_heading_level, docx_document
                    )
                    if current_chapter is not None:
                        current_chapter.elements.extend(subchapters)
                        logger.debug(f"Added {len(subchapters)} subchapters to Chapter '{current_chapter.title}'.")
            elif isinstance(block, DocxTable):
                if current_chapter is not None:
                    table = Table.from_word(block)
                    current_chapter.elements.append(table)
                    logger.debug(f"Added Table to Chapter '{current_chapter.title}'.")
                current_index += 1
            elif isinstance(block, DocxParagraph):
                if block.text.strip():
                    if current_chapter is not None:
                        paragraph = Paragraph.from_word(block)
                        current_chapter.elements.append(paragraph)
                        logger.debug(f"Added Paragraph to Chapter '{current_chapter.title}'.")
                current_index += 1
            else:
                logger.warning(f"Unknown block type at index {current_index}: {type(block).__name__}")
                current_index += 1

        if current_chapter is not None:
            chapters.append(current_chapter)
            logger.debug(f"Added final Chapter '{current_chapter.title}' to chapters list.")

        logger.info(f"Parsed {len(chapters)} chapters starting from index {current_index}.")
        return chapters, current_index

    def to_word(self, docx_document: DocxDocument, level: int = 1) -> None:
        """
        Add the chapter and its elements to a Word document.
        """
        logger.info(f"Adding Chapter '{self.title}' to Word document at level {level}.")
        docx_document.add_heading(self.title, level=level)
        for element in self.elements:
            element.to_word(docx_document, level=level + 1)
        logger.debug(f"Chapter '{self.title}' added to Word document successfully.")

    def to_xml(self) -> ET.Element:
        """
        Convert the chapter to an XML element.
        """
        logger.info(f"Converting Chapter '{self.title}' to XML.")
        chapter_element = ET.Element('chapter', title=self.title, id=self.id_)
        for element in self.elements:
            chapter_element.append(element.to_xml())
        logger.debug(f"Chapter '{self.title}' converted to XML successfully.")
        return chapter_element

    def __repr__(self) -> str:
        return f"Chapter(title='{self.title}', id='{self.id_}', elements={self.elements})"


class Document:
    """
    Represents the entire document, containing a title and a list of elements.
    """

    def __init__(self, title: str, elements: List[DocumentElement]):
        self.title = title.strip()
        self.elements = elements
        logger.debug(f"Initialized Document with title: '{self.title}' and {len(self.elements)} elements.")

    @classmethod
    def from_xml(cls, xml_string: str) -> 'Document':
        """
        Create a Document instance from an XML string.
        """
        logger.info("Parsing Document from XML string.")
        root = ET.fromstring(xml_string)
        title = root.attrib.get('title', 'Untitled Document')
        elements: List[DocumentElement] = []

        for chapter_element in root.findall('chapter'):
            elements.append(Chapter.from_xml(chapter_element))

        logger.debug(f"Parsed Document with title: '{title}' and {len(elements)} chapters.")
        return cls(title, elements)

    @classmethod
    def from_word(cls, file_path: str) -> 'Document':
        """
        Create a Document instance from a Word (.docx) file.
        """
        logger.info(f"Parsing Document from Word file: '{file_path}'.")
        docx_document = DocxDocument(file_path)
        blocks = list(iter_block_items(docx_document))  # Collect all paragraphs and tables

        chapters: List[Chapter] = []
        index = 0
        title: Optional[str] = None

        while index < len(blocks):
            block = blocks[index]
            if isinstance(block, DocxParagraph) and block.style.name.startswith('Title'):
                title = block.text.strip()
                logger.debug(f"Document title found: '{title}'.")
                index += 1
            else:
                new_chapters, index = Chapter.from_word(blocks, index, 1, docx_document)
                chapters.extend(new_chapters)

        if not title:
            title = "Untitled Document"
            logger.warning("No title found in the Word document. Using default title.")

        logger.info(f"Parsed Document titled '{title}' with {len(chapters)} chapters.")
        return cls(title, chapters)

    def to_word(self, file_name: str) -> None:
        """
        Save the Document to a Word (.docx) file.
        """
        logger.info(f"Saving Document to Word file: '{file_name}'.")
        docx_document = DocxDocument()
        docx_document.add_heading(self.title, level=0)
        logger.debug(f"Added document title '{self.title}' to Word document.")

        for element in self.elements:
            element.to_word(docx_document)

        docx_document.save(file_name)
        logger.info(f"Document saved to '{file_name}' successfully.")

    def to_xml(self) -> str:
        """
        Convert the Document to an XML string.
        """
        logger.info("Converting Document to XML string.")
        document_element = ET.Element('document', title=self.title)
        for element in self.elements:
            document_element.append(element.to_xml())
        xml_string = ET.tostring(document_element, encoding='unicode')
        logger.debug("Document converted to XML successfully.")
        return xml_string

    def __repr__(self) -> str:
        return f"Document(title='{self.title}', elements={self.elements})"
